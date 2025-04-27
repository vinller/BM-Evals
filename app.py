from flask import Flask, render_template, request, redirect, url_for, session, send_file, jsonify, send_from_directory
import os
import json
import random
import uuid
import zipfile
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from io import BytesIO

app = Flask(__name__)
app.secret_key = 'supersecretkey'

# Load user data
with open("users.json") as f:
    USERS = json.load(f)

QUESTIONS_FILE = "questions_grouped.json"

# --- Helper Functions ---
def get_current_questions():
    with open(QUESTIONS_FILE) as f:
        return json.load(f)

# --- Routes ---
@app.route("/History/<path:filename>")
def download_file(filename):
    return send_from_directory("History", filename, as_attachment=True)

@app.route("/")
def home():
    if "user" in session:
        return redirect("/dashboard")
    return redirect("/login")

@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        email = request.form.get("email")
        password = request.form.get("password")
        
        for user_id, info in USERS.items():
            # Check if email and password match
            if info["email"] == email and info["password"] == password:
                session["user"] = info
                
                # Redirect to respective dashboard based on role
                if info["role"] == "admin":
                    return redirect("/dashboard")
                elif info["role"] == "candidate":
                    return redirect("/candidate_dashboard")  # Redirect to take evaluation for candidates
                
        return render_template("login.html", error="Invalid credentials")
    return render_template("login.html")

@app.route("/logout")
def logout():
    session.pop("user", None)
    return redirect("/login")

@app.route("/dashboard")
def dashboard():
    if "user" not in session:
        return redirect("/login")
    return render_template("dashboard.html", user=session["user"])

@app.route("/generate", methods=["GET", "POST"])
def generate():
    if "user" not in session:
        return redirect("/login")

    if request.method == "POST":
        score_threshold = request.form.get("score_threshold")

        if not score_threshold:
            score_threshold = 80
        else:
            score_threshold = int(score_threshold)

        count = int(request.form.get("set_count"))
        student_copy = request.form.get("student_copy", "yes")
        filename_format = request.form.get("filename_format", "month")
        selected_sections = request.form.getlist("sections")
        eval_title = request.form.get("evaluation_title") or ""

        # Handle default all sections if none selected
        if not selected_sections:
            selected_sections = [str(i) for i in range(1, 11)]

        timestamp = datetime.now()
        month_name = timestamp.strftime("%B")
        date_str = timestamp.strftime("%Y-%m-%d")
        year = timestamp.strftime("%Y")

        if filename_format == "month":
            prefix = f"{month_name} {year}"
        elif filename_format == "term":
            prefix = f"Spring {year}"  # You could make it dynamic if you want
        else:
            prefix = date_str

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, 'w') as zipf:
            for _ in range(count):
                eval_id = ''.join(random.choices('ABCDEFGHIJKLMNOPQRSTUVWXYZ0123456789', k=6))
                if student_copy == "yes":
                    student_doc = generate_doc(eval_id, "Student Copy", session["user"], prefix, selected_sections, eval_title)
                    student_name = f"{eval_id} {prefix} Student Copy.docx"
                    student_path = os.path.join("History", student_name)
                    student_doc.save(student_path)
                    zipf.write(student_path, arcname=student_name)

                evaluator_doc = generate_doc(eval_id, "Evaluators Copy", session["user"], prefix, selected_sections, eval_title, include_answers=True)
                evaluator_name = f"{eval_id} {prefix} Evaluators Copy.docx"
                evaluator_path = os.path.join("History", evaluator_name)
                evaluator_doc.save(evaluator_path)
                zipf.write(evaluator_path, arcname=evaluator_name)

        zip_buffer.seek(0)
        zip_filename = f"{prefix} {count} Set{'s' if count > 1 else ''}.zip"
        
        zip_path = os.path.join("History", zip_filename)
        with open(zip_path, "wb") as f:
            f.write(zip_buffer.read())

# Redirect to the download ready page
        return redirect(url_for('download_ready', filename=zip_filename))
    

    return render_template("generate.html", user=session["user"])

def generate_doc(uuid, copy_type, user, prefix, selected_sections, eval_title, include_answers=False, score_threshold=80):
    doc = Document()
    evaluation_data = {
    "uuid": uuid,
    "title": eval_title if eval_title else f"BM Evaluation {prefix}",
    "threshold": score_threshold,
    "questions": []
}

    # Heading
    if eval_title:
        doc.add_heading(eval_title, 0)
    else:
        doc.add_heading(f"BM Evaluation {prefix}", 0)

    # Name/Score/Date
    header = doc.add_paragraph()
    header.add_run("Name: __________________________   ").bold = False
    header.add_run("Date: ________________").bold = False
    header.add_run("Score: ________/100   ").bold = False

    # ID and By
    id_by = doc.add_paragraph()
    id_by.add_run(f"Evaluation ID: {uuid}    ").bold = False
    id_by.add_run(f"Generated by: {user['first_name']} {user['last_name']}").bold = False

    # Instructions
    if copy_type == "Student Copy":
        doc.add_paragraph("Instructions:")
        instructions = [
            "This evaluation consists of written questions. Answer clearly and concisely.",
            "You must complete the written portion without using phones, notes, documents, or external sources.",
            "Attempt every question. Partial credit is awarded.",
            f"You must score {score_threshold}% or more to pass."
        ]
        table = doc.add_table(rows=1, cols=1)
        cell = table.rows[0].cells[0]
        for line in instructions:
            cell.add_paragraph(f"• {line}")
    else:
        doc.add_paragraph("Instructions: Use the provided answers to guide grading. Each question is worth 2 points.")

    random.seed(uuid)
    current_questions = get_current_questions()

    q_number = 1
    for section_name, questions in current_questions.items():
        section_num = section_name.split(" ")[1]
        if section_num not in selected_sections:
            continue

        doc.add_paragraph(f"\n{section_name}", style="Heading 2")
        sample_count = min(10, len(questions))
        sampled = random.sample(questions, sample_count)

        for q in sampled:
            lines = q["q"].split("\n")
            p = doc.add_paragraph()
            p.add_run(f"Q{q_number}. {lines[0]}").bold = True
            for line in lines[1:]:
                p.add_run(f"\n    {line}")

            # Save question and answer for the JSON
            evaluation_data["questions"].append({
    "q": q["q"],
    "a": q["a"]
})

            if copy_type == "Student Copy":
                doc.add_paragraph("\n\n\n\n")
            elif include_answers:
                answer_lines = q["a"].split("\n")
                answer_para = doc.add_paragraph("Answer: ")
                answer_para.add_run(answer_lines[0])
                for line in answer_lines[1:]:
                    answer_para.add_run(f"\n    {line}")

            q_number += 1

    # ONLY save the evaluation JSON once per evaluation
    if copy_type == "Student Copy":  # Only once
        os.makedirs("Evaluations", exist_ok=True)
    with open(f"Evaluations/{uuid}.json", "w") as f:
        json.dump(evaluation_data, f, indent=2)

    return doc

@app.route("/download_ready/<filename>")
def download_ready(filename):
    return render_template("download_ready.html", filename=filename, user=session["user"])

@app.route("/score/<uuid>")
def start_scoring(uuid):
    if "user" not in session:
        return redirect("/login")

    try:
        with open(f"Evaluations/{uuid}.json") as f:
            eval_data = json.load(f)
    except FileNotFoundError:
        return "Evaluation not found", 404

    return render_template("score_evaluation.html",
        uuid=uuid,
        questions=eval_data["questions"],
        threshold=eval_data["threshold"],
        user=session["user"]
    )


@app.route("/lookup", methods=["GET", "POST"])
def lookup():
    if "user" not in session:
        return redirect("/")

    files = []
    if request.method == "POST":
        uuid = request.form.get("uuid")
        if uuid:
            for filename in os.listdir("History"):
                if uuid.upper() in filename.upper():
                    files.append(filename)
    return render_template("lookup.html", user=session["user"], files=files)

@app.route("/questions")
def question_bank():
    if "user" not in session:
        return redirect("/login")
    
    with open(QUESTIONS_FILE) as f:
        grouped = json.load(f)

    # Pass both section numbers and section names
    sections = []
    for idx, section_name in enumerate(grouped.keys(), start=1):
        sections.append((idx, section_name))
    
    return render_template("questions.html", user=session["user"], sections=sections)


@app.route("/questions/section_<int:section_num>")
def view_section(section_num):
    if "user" not in session:
        return redirect("/login")

    with open(QUESTIONS_FILE) as f:
        grouped = json.load(f)

    section_keys = list(grouped.keys())
    section_key = section_keys[section_num - 1]
    questions = grouped[section_key]

    return render_template("section_questions.html", section_num=section_num, questions=questions, section_name=section_key, user=session["user"])

@app.route("/questions/section_<int:section_num>/delete", methods=["POST"])
def delete_questions(section_num):
    to_delete = request.json.get("questions", [])

    with open(QUESTIONS_FILE) as f:
        grouped = json.load(f)

    section_keys = list(grouped.keys())
    section_key = section_keys[section_num - 1]
    grouped[section_key] = [q for q in grouped[section_key] if q["q"] not in to_delete]

    with open(QUESTIONS_FILE, "w") as f:
        json.dump(grouped, f, indent=2)

    return jsonify({"status": "success"})

@app.route("/questions/section_<int:section_num>/add", methods=["POST"])
def add_question(section_num):
    data = request.json
    new_q = {"q": data["q"], "a": data["a"]}

    with open(QUESTIONS_FILE) as f:
        grouped = json.load(f)

    section_keys = list(grouped.keys())
    section_key = section_keys[section_num - 1]
    grouped[section_key].append(new_q)

    with open(QUESTIONS_FILE, "w") as f:
        json.dump(grouped, f, indent=2)

    return jsonify({"status": "success"})

@app.route("/questions/section_<int:section_num>/edit", methods=["POST"])
def edit_question(section_num):
    data = request.json
    old_q = data["old_q"]
    new_q = data["new_q"]
    new_a = data["new_a"]

    with open(QUESTIONS_FILE) as f:
        grouped = json.load(f)

    section_keys = list(grouped.keys())
    section_key = section_keys[section_num - 1]

    # Find and update the question
    for question in grouped[section_key]:
        if question["q"] == old_q:
            question["q"] = new_q
            question["a"] = new_a
            break

    with open(QUESTIONS_FILE, "w") as f:
        json.dump(grouped, f, indent=2)

    return jsonify({"status": "success"})

@app.route("/results", methods=["GET", "POST"])
def results_page():
    if "user" not in session:
        return redirect("/login")
    
    if request.method == "POST":
        uuid = request.form.get("uuid").upper()
        filepath = os.path.join("Evaluations", f"{uuid}.json")

        if not os.path.exists(filepath):
            return "Evaluation not found!", 404
        
        with open(filepath) as f:
            eval_data = json.load(f)
        
        return render_template("score_evaluation.html", uuid=uuid, questions=eval_data["questions"], threshold=eval_data["threshold"], user=session["user"])
    
    return render_template("results.html", user=session["user"])

@app.route("/results/save", methods=["POST"])
def save_result():
    if "user" not in session:
        return redirect("/login")
    
    # Collect form data
    uuid = request.form.get("uuid")
    threshold = int(request.form.get("threshold", 80))
    candidate_name = request.form.get("candidate_name")  # Candidate's name from the form
    date_taken = request.form.get("date_taken")  # Date when the evaluation was taken
    evaluator = f"{session['user']['first_name']} {session['user']['last_name']}"  # Evaluator's name
    
    # Initialize variables to calculate score and details
    scores = []
    total_score = 0
    details = []  # This will hold 'c' for correct, 'p' for partial, 'w' for wrong

    # Loop through the scores from the form and calculate total score and details
    for key in request.form.keys():
        if key.startswith("score_"):
            score = int(request.form.get(key))  # Extract score
            scores.append(score)
            total_score += score
            if score == 2:
                details.append("c")  # Correct
            elif score == 1:
                details.append("p")  # Partial credit
            else:
                details.append("w")  # Wrong

    # Calculate the total percentage score
    percent = (total_score / (len(scores) * 2)) * 100
    status = "PASS" if percent >= threshold else "FAIL"  # Determine if it's a pass or fail

    # File to store results
    results_file = "results.json"
    
    # Create the file if it doesn't exist
    if not os.path.exists(results_file):
        with open(results_file, "w") as f:
            json.dump([], f, indent=2)
    
    # Read existing results
    with open(results_file) as f:
        data = json.load(f)

    # Add the new result to the list
    data.append({
        "uuid": uuid,
        "candidate_name": candidate_name,
        "date": date_taken,
        "evaluator": evaluator,
        "score": round(percent),  # Rounded score
        "threshold": threshold,
        "status": status,
        "details": details,  # Store details about each question (correct/partial/wrong)
        "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S")  # Record the time the evaluation was saved
    })

    # Write the updated data back to the results file
    with open(results_file, "w") as f:
        json.dump(data, f, indent=2)

    # Pass the status to template for the screen effect
    return render_template("popup_result.html", status=status)

@app.route("/results/check/<uuid>")
def check_result(uuid):
    if not os.path.exists("results.json"):
        return jsonify({"exists": False})

    with open("results.json") as f:
        results = json.load(f)

    exists = any(entry["uuid"].upper() == uuid.upper() for entry in results)
    return jsonify({"exists": exists})

@app.route("/score_lookup", methods=["GET", "POST"])
def score_lookup():
    if "user" not in session:
        return redirect("/login")

    candidates = []
    if os.path.exists("results.json"):
        with open("results.json") as f:
            results = json.load(f)
            candidates = list({r["candidate_name"] for r in results})

    if request.method == "POST":
        uuid = request.form.get("uuid", "").strip().upper()
        candidate = request.form.get("candidate", "").strip()

        if uuid:
            return redirect(url_for("candidate_profile", uuid=uuid))
        elif candidate:
            with open("results.json") as f:
                results = json.load(f)
                for r in results:
                    if r["candidate_name"].lower() == candidate.lower():
                        return redirect(url_for("candidate_profile", uuid=r["uuid"]))

    return render_template("score_lookup.html", candidates=candidates, user=session["user"])
@app.route("/candidate/<uuid>")
def candidate_profile(uuid):
    if "user" not in session:
        return redirect("/login")
    try:
        # Load results.json
        with open("results.json") as f:
            results = json.load(f)

        # Find the specific evaluation record
        record = next((r for r in results if r["uuid"].upper() == uuid.upper()), None)
        if not record:
            return "Evaluation not yet graded.", 404

        # Find matching evaluator DOCX
        matching_docx = None
        history_folder = "History"
        if os.path.exists(history_folder):
            for file in os.listdir(history_folder):
                if uuid.upper() in file and "Evaluator" in file:
                    matching_docx = file
                    break

        if not matching_docx:
            matching_docx = "None found"

        # Load the evaluation questions
        eval_file_path = f"Evaluations/{uuid.upper()}.json"
        if not os.path.exists(eval_file_path):
            return "Evaluation file not found.", 404

        with open(eval_file_path) as f:
            eval_data = json.load(f)

        # Find all past evaluations of this candidate
        candidate_name = record["candidate_name"]
        # Load all past evaluations by the candidate
        past_evals = []
        with open("results.json") as f:
            all_results = json.load(f)
        for r in all_results:
            if r["candidate_name"].lower() == record["candidate_name"].lower():
                past_evals.append({
                    "uuid": r["uuid"],
                    "status": r["status"],
                    "score": r["score"],
                    "date": r.get("date", "N/A")
            })

        return render_template(
            "candidate_profile.html",
            record=record,
            questions=eval_data["questions"],
            threshold=eval_data["threshold"],
            docx_filename=matching_docx,
            uuid=uuid,
            user=session["user"],
            enumerate=enumerate,
            past_evals=past_evals
        )

    except Exception as e:
        print(f"Error loading candidate profile: {e}")
    return "Error loading candidate profile.", 500

@app.route("/lookup_score", methods=["POST"])
def lookup_score():
    uuid = request.form.get("uuid")
    candidate = request.form.get("candidate")

    # Priority: if UUID provided, use that
    if uuid:
        return redirect(url_for("candidate_profile", uuid=uuid.upper()))
    elif candidate:
        # Look up UUID from candidate name in results.json
        try:
            with open("results.json") as f:
                results = json.load(f)
            for record in results:
                if record["candidate_name"].lower() == candidate.lower():
                    return redirect(url_for("candidate_profile", uuid=record["uuid"]))
        except Exception as e:
            print(e)
            return "Error finding candidate", 500

    return "Invalid lookup", 400

def fix_results_file():
    results_file = "results.json"

    try:
        with open(results_file) as f:
            results = json.load(f)
        
        updated = False

        for entry in results:
            # If 'details' field missing, add a placeholder
            if "details" not in entry:
                num_questions = 10  # Adjust based on how many questions your evaluations usually have
                entry["details"] = ["c"] * num_questions  # Assume full correct if missing
                updated = True

        if updated:
            with open(results_file, "w") as f:
                json.dump(results, f, indent=2)
            print("✅ Fixed missing 'details' field for old evaluations.")
        else:
            print("✅ No fixes needed. All evaluations already have 'details'.")

    except Exception as e:
        print(f"❌ Error: {e}")

@app.route("/take_evaluation/uuid", methods=["POST"])
def take_evaluation_uuid():
    if "user" not in session or session["user"]["role"] != "candidate":
        return redirect("/login")  # Ensure the user is logged in and is a candidate

    uuid = request.form.get("uuid").strip().upper()

    # Check if the evaluation exists for this UUID
    eval_path = os.path.join("Evaluations", f"{uuid}.json")
    if not os.path.exists(eval_path):
        return "Evaluation not found", 404

    # Load the evaluation data
    with open(eval_path) as f:
        eval_data = json.load(f)

    # Start the countdown timer (set for 2 hours in seconds)
    countdown = 2 * 60 * 60  # 2 hours in seconds

    return render_template(
        "evaluation_questions.html",
        uuid=uuid,
        questions=eval_data["questions"],
        threshold=eval_data["threshold"],
        countdown=countdown,
        user=session["user"]
    )

@app.route("/candidate_dashboard")
def candidate_dashboard():
    if "user" not in session or session["user"]["role"] != "candidate":
        return redirect("/login")  # Redirect to login if not logged in as a candidate
    
    # Render the candidate dashboard page
    return render_template("candidate_dashboard.html", user=session["user"])
from flask import Flask
from datetime import datetime

# Create the custom date format filter
def datetimeformat(value, format='%B %d, %Y'):
    if value:
        return datetime.strptime(value, '%Y-%m-%d').strftime(format)
    return value

app.jinja_env.filters['datetimeformat'] = datetimeformat


if __name__ == "__main__":
    os.makedirs("History", exist_ok=True)
    app.run(debug=True)
